#!/usr/bin/env python3
"""Convert Basera graduation study Markdown to academic Arabic RTL Word (.docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_FONT = "Traditional Arabic"
FALLBACK_FONT = "Arial"
HEADING_COLOR = RGBColor(0x1B, 0x4F, 0x72)
MUTED = RGBColor(0x5A, 0x6A, 0x7A)


def set_paragraph_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_run_rtl(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    r_pr.append(rtl)


def set_run_font(run, size: Pt, *, bold: bool = False, name: str = BODY_FONT) -> None:
    run.bold = bold
    run.font.name = name
    run.font.size = size
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), FALLBACK_FONT)
    r_fonts.set(qn("w:hAnsi"), FALLBACK_FONT)
    r_fonts.set(qn("w:cs"), name)
    set_run_rtl(run)


def shade_cell(cell, fill: str = "D6EAF8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)
    run._r.append(end)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]

    hp = section.header.paragraphs[0]
    set_paragraph_rtl(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("دراسة مشروع التخرج — نظام بصيرة (Basera)")
    set_run_font(run, Pt(10), name=FALLBACK_FONT)
    run.font.color.rgb = MUTED

    fp = section.footer.paragraphs[0]
    set_paragraph_rtl(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("صفحة ")
    set_run_font(run, Pt(10), name=FALLBACK_FONT)
    add_page_number(fp)
    run2 = fp.add_run(" — مشروع بصيرة")
    set_run_font(run2, Pt(10), name=FALLBACK_FONT)


def add_formatted_text(
    paragraph,
    text: str,
    *,
    bold_default: bool = False,
    size: Pt = Pt(14),
) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size, bold=bold_default)
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            set_run_font(run, size, bold=True)
        else:
            run = paragraph.add_run(chunk[1:-1])
            set_run_font(run, Pt(11), name="Consolas")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size, bold=bold_default)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def is_separator_row(line: str) -> bool:
    s = line.strip().strip("|").strip()
    if not s:
        return False
    return all(re.match(r"^:?-+:?$", cell.strip()) for cell in s.split("|"))


def parse_table_rows(lines: list[str]) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = 0
    while i < len(lines) and is_table_row(lines[i]):
        if not is_separator_row(lines[i]):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_rtl(p)
            add_formatted_text(p, cell_text, size=Pt(12))
            for run in p.runs:
                if ri == 0:
                    run.bold = True
            if ri == 0:
                shade_cell(cell)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int, *, is_first_title: bool = False) -> None:
    clean = text.lstrip("#").strip()
    sizes = {1: 26, 2: 18, 3: 16, 4: 14, 5: 13}

    # فاصل صفحة قبل الأبواب الرئيسية المرقّمة
    if level == 2 and re.match(r"^#{2}\s+\d", text):
        doc.add_page_break()

    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_first_title else WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(16 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(clean)
    set_run_font(run, Pt(sizes.get(level, 12)), bold=True, name=FALLBACK_FONT)
    run.font.color.rgb = HEADING_COLOR


def add_codeblock(doc: Document, lines: list[str], language: str = "") -> None:
    label = (
        "مخطط (Mermaid) — يُعرض في محرر Markdown أو VS Code، ويُنصح بتصديره صورة للطباعة:"
        if language == "mermaid"
        else "مقتطف برمجي:"
    )
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    run = p.add_run(label)
    set_run_font(run, Pt(11), bold=True, name=FALLBACK_FONT)
    run.italic = True
    run.font.color.rgb = MUTED

    # اختصار مخططات Mermaid الطويلة في Word
    display = lines
    if language == "mermaid" and len(lines) > 24:
        display = lines[:20] + ["… (المخطط الكامل في ملف Markdown)"]

    for line in display:
        cp = doc.add_paragraph()
        set_paragraph_rtl(cp)
        cp.paragraph_format.left_indent = Cm(0.4)
        cp.paragraph_format.space_after = Pt(0)
        run = cp.add_run(line)
        set_run_font(run, Pt(9), name="Consolas")
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.add_paragraph()


def style_body_paragraph(p) -> None:
    set_paragraph_rtl(p)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles["Normal"]
    styles.font.name = BODY_FONT
    styles.font.size = Pt(14)
    pf = styles.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    add_header_footer(doc)


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_document(doc)

    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    first_h1 = True

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip().lower()
                code_lines = []
            else:
                add_codeblock(doc, code_lines, code_lang)
                in_code = False
                code_lang = ""
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        m = re.match(r"^(#{1,5})\s+(.+)$", line)
        if m:
            lvl = len(m.group(1))
            add_heading(doc, line, lvl, is_first_title=(first_h1 and lvl == 1))
            if first_h1 and lvl == 1:
                first_h1 = False
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            set_paragraph_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 28)
            set_run_font(run, Pt(11), name=FALLBACK_FONT)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            style_body_paragraph(p)
            add_formatted_text(p, stripped.lstrip(">").strip())
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = HEADING_COLOR
            i += 1
            continue

        if is_table_row(line):
            rows, consumed = parse_table_rows(lines[i:])
            add_table(doc, rows)
            i += consumed
            continue

        m = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m:
            p = doc.add_paragraph()
            style_body_paragraph(p)
            p.paragraph_format.right_indent = Cm(0.4)
            add_formatted_text(p, f"{m.group(1)}. {m.group(2)}")
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph()
            style_body_paragraph(p)
            p.paragraph_format.right_indent = Cm(0.4)
            add_formatted_text(p, "• " + stripped[2:])
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        style_body_paragraph(p)
        add_formatted_text(p, stripped)
        i += 1

    doc.save(str(docx_path))
    print(f"Saved OK ({docx_path.stat().st_size // 1024} KB)")


def main() -> int:
    base = Path(__file__).resolve().parent.parent
    md_path = base / "دراسة مشروع التخرج.md"
    docx_path = base / "دراسة مشروع التخرج.docx"

    if len(sys.argv) >= 2:
        md_path = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        docx_path = Path(sys.argv[2])

    if not md_path.exists():
        print(f"Missing markdown", file=sys.stderr)
        return 1

    convert_md_to_docx(md_path, docx_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
